"""Step 1 of the RAG pipeline: ingest source documents.

Loads supported files (PDF, DOCX, XLSX, TXT, MD, CSV, ...) — a single file or
a whole folder — splits them into overlapping text chunks, embeds the chunks
with OllamaEmbeddings, and persists them in a local Chroma vector store.

Features:
  * --source <file-or-folder>  ingest one file or all supported files in a
                               folder. Defaults to data/pdfs/ (SOURCE_DEFAULT).
  * default: APPEND            new files are added to whatever is already in
                               the store, so a corpus can grow across
                               sessions (Unit 1 + 2 today, Unit 3 tomorrow).
  * --clear                    wipe the existing collection before ingesting
                               the new source (for switching corpora).
  * --list                     print which source files are currently in the
                               store, without ingesting anything.

Duplicate guard: a file already in the store (matched by name + content hash)
is skipped with a warning on append, so re-runs don't duplicate chunks.

Usage:
    python -m src.ingest                        # append data/pdfs
    python -m src.ingest --source data/unit3     # append a folder
    python -m src.ingest --source data/unit3.pdf # append a single file
    python -m src.ingest --clear --source new_subject
    python -m src.ingest --list
    python -m src.ingest --analyze
"""
from __future__ import annotations

import hashlib
import sys
from pathlib import Path

import chromadb
from langchain_core.documents import Document
from langchain_community.vectorstores import Chroma
from langchain_ollama import OllamaEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

from src.config import (
    CHUNK_OVERLAP,
    CHUNK_SIZE,
    COLLECTION_NAME,
    OLLAMA_EMBED_MODEL,
    OLLAMA_HOST,
    SOURCE_DEFAULT,
    VECTORSTORE_DIR,
)

# File extensions the ingestion pipeline can parse into text for embedding.
# Kept intentionally lightweight (no fragile `unstructured` dependency):
#   pdf   -> pypdf        (already a project dependency)
#   docx  -> python-docx   (added dep)
#   xlsx/xls -> openpyxl   (added dep)
#   txt/md/csv/log/json -> plain text readers (stdlib)
SUPPORTED_EXTS = {
    ".pdf",
    ".docx",
    ".txt",
    ".md",
    ".csv",
    ".xlsx",
    ".xls",
    ".log",
    ".json",
}


def resolve_source_files(source: Path | str) -> list[Path]:
    """Resolve --source into a list of supported files, validating existence.

    Raises FileNotFoundError with a clear message if the path is missing or
    holds no supported files.
    """
    p = Path(source)

    if not p.exists():
        raise FileNotFoundError(f"Source path does not exist: {p}")

    if p.is_file():
        if p.suffix.lower() not in SUPPORTED_EXTS:
            ext = p.suffix.lower() or "(none)"
            raise FileNotFoundError(
                f"Unsupported file type '{ext}'. Supported: {sorted(SUPPORTED_EXTS)}"
            )
        files = [p]
    else:
        files = sorted(
            fn for fn in p.iterdir() if fn.is_file() and fn.suffix.lower() in SUPPORTED_EXTS
        )

    if not files:
        raise FileNotFoundError(f"No supported files found at: {p}")
    return files


# Kept as an alias so existing callers/tests that referenced the old name
# still work. It now accepts any supported file type, not just PDFs.
resolve_pdf_files = resolve_source_files


def _get_client() -> chromadb.ClientAPI:
    return chromadb.PersistentClient(path=str(VECTORSTORE_DIR))


def _existing_keys() -> set[tuple[str, str]]:
    """Return {(source_name, file_sha1)} already present in the store."""
    client = _get_client()
    try:
        col = client.get_collection(COLLECTION_NAME)
        data = col.get(include=["metadatas"])
    except Exception:
        return set()
    metas = data.get("metadatas") or []
    return {
        (m["source"], m["file_sha1"]) for m in metas
        if m and m.get("source") and m.get("file_sha1")
    }


def list_sources() -> list[str]:
    """Return the distinct source filenames currently in the store."""
    client = _get_client()
    try:
        col = client.get_or_create_collection(COLLECTION_NAME)
        data = col.get(include=["metadatas"])
    except Exception:
        return []
    metas = data.get("metadatas") or []
    sources = sorted({m.get("source") for m in metas if m and m.get("source")})
    return sources


def _file_sha1(path: Path) -> str:
    """Streaming SHA-1 of a file, used for duplicate detection."""
    h = hashlib.sha1()
    with open(path, "rb") as fh:
        for block in iter(lambda: fh.read(65536), b""):
            h.update(block)
    return h.hexdigest()


def _read_as_documents(path: Path) -> list[Document]:
    """Parse a single file into LangChain documents based on its extension."""
    ext = path.suffix.lower()

    if ext == ".pdf":
        from langchain_community.document_loaders import PyPDFLoader

        return PyPDFLoader(str(path)).load()

    if ext == ".docx":
        import docx

        text = "\n".join(p.text for p in docx.Document(str(path)).paragraphs if p.text)
        return [Document(page_content=text)]

    if ext in (".xlsx", ".xls"):
        import openpyxl

        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        rows: list[str] = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None and str(c).strip()]
                if cells:
                    rows.append("\t".join(cells))
        wb.close()
        return [Document(page_content="\n".join(rows))]

    # Plain-text formats (txt, md, csv, log, json). Read with encoding fallback.
    try:
        content = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        content = path.read_text(encoding="latin-1")
    return [Document(page_content=content)]


def load_documents(files: list[Path]) -> list:
    """Load supported files into LangChain documents.

    Tags each page/document with source filename and a content hash (file_sha1)
    used for duplicate detection. Citation metadata (source, page) is unchanged
    for PDFs; other formats get page=0.
    """
    documents = []
    for src in files:
        print(f"Loading {src.name} ...")
        digest = _file_sha1(src)
        try:
            pages = _read_as_documents(src)
        except Exception as exc:
            print(f"  ERROR loading {src.name}: {exc}")
            continue
        for page in pages:
            page.metadata["source"] = src.name
            page.metadata["file_sha1"] = digest
            page.metadata.setdefault("page", 0)
        documents.extend(pages)
    return documents


# Kept as an alias for any existing callers/tests referencing the old name.
load_pdfs = load_documents


def split_documents(documents: list) -> tuple:
    """Split loaded page documents into chunks.

    Returns (chunks, per_page_over_CHUNK_SIZE) where the second value is the
    count of pages that exceeded CHUNK_SIZE and had to be split.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
    )
    chunks = splitter.split_documents(documents)
    oversized = sum(1 for d in documents if len(d.page_content) > CHUNK_SIZE)
    return chunks, oversized


def _embed_chunks(chunks: list) -> None:
    """Embed and append chunks into the persistent Chroma store (default: append)."""
    embeddings = OllamaEmbeddings(
        base_url=OLLAMA_HOST,
        model=OLLAMA_EMBED_MODEL,
    )
    vectorstore = Chroma(
        collection_name=COLLECTION_NAME,
        persist_directory=str(VECTORSTORE_DIR),
        embedding_function=embeddings,
    )
    if chunks:
        vectorstore.add_documents(chunks)


def clear_store() -> None:
    """Wipe the existing Chroma collection entirely."""
    client = _get_client()
    try:
        client.delete_collection(COLLECTION_NAME)
        print(f"Cleared existing collection '{COLLECTION_NAME}'.")
    except Exception:
        print(f"No existing collection '{COLLECTION_NAME}' to clear.")


def ingest(source: Path | str, clear: bool = False) -> int:
    """Resolve, filter, and ingest supported source files; return count added."""
    # Validate the source FIRST so that a bad --source can never trigger the
    # destructive --clear wipe (resolve_source_files raises before clear_store).
    files = resolve_source_files(source)

    if clear:
        clear_store()
        pending = files
        dupes = []
    else:
        existing = _existing_keys()
        pending, dupes = [], []
        for f in files:
            if (f.name, _file_sha1(f)) in existing:
                dupes.append(f)
            else:
                pending.append(f)

    if not pending:
        for f in dupes:
            print(f"  Skipped: {f.name} (already in store, name + content-hash match)")
        print(f"Nothing new to ingest from {source}.")
        return 0

    documents = load_documents(pending)
    chunks, oversized = split_documents(documents)
    pages = len(documents)
    print(f"Pages: {pages} | Chunks: {len(chunks)} | "
          f"Slides over {CHUNK_SIZE} chars split: {oversized}")
    print(f"Ratio: ~{len(chunks) / max(pages, 1):.2f} chunks per page")

    _embed_chunks(chunks)
    print(f"Added {pages} pages / {len(chunks)} chunks from {len(pending)} file(s) "
          f"to '{COLLECTION_NAME}' in {VECTORSTORE_DIR}")

    for f in dupes:
        print(f"  Skipped: {f.name} (already in store, name + content-hash match)")
    return len(pending)


def analyze(source: Path | str) -> None:
    """Split-only diagnostic: report chunking without embedding."""
    files = resolve_source_files(source)
    documents = load_documents(files)
    if not documents:
        return

    chunks, oversized = split_documents(documents)
    print(f"\nPages: {len(documents)} / {len(files)} file(s)")
    print(f"Chunks: {len(chunks)}  (oversized slides split: {oversized})")
    print(f"Ratio: ~{len(chunks) / max(len(documents), 1):.2f} chunks per page\n")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len,
        separators=["\n\n", "\n", ". ", " "],
    )
    for d in documents:
        if len(d.page_content) > CHUNK_SIZE:
            sub = splitter.split_text(d.page_content)
            print(f"  Split slide: {d.metadata.get('source')} page {d.metadata.get('page')} "
                  f"({len(d.page_content)} chars -> {len(sub)} sub-chunks)")


def _cli() -> None:
    import argparse

    parser = argparse.ArgumentParser(
        prog="python -m src.ingest",
        description="Ingest course PDFs into the Chroma vector store.",
    )
    parser.add_argument(
        "--source",
        default=str(SOURCE_DEFAULT),
        help=f"PDF file or folder to ingest. Default: {SOURCE_DEFAULT}",
    )
    parser.add_argument(
        "--clear",
        action="store_true",
        help="Wipe the existing collection before ingesting (full reset).",
    )
    parser.add_argument(
        "--list",
        action="store_true",
        help="List source files currently in the store, then exit.",
    )
    parser.add_argument(
        "--analyze",
        action="store_true",
        help="Only split + report chunking (no embedding, no store changes).",
    )
    args = parser.parse_args()

    try:
        if args.list:
            sources = list_sources()
            if sources:
                print("Source files currently in the vector store:")
                for s in sources:
                    print(f"  - {s}")
            else:
                print("The vector store is empty.")
            return
        if args.analyze:
            analyze(args.source)
        else:
            ingest(args.source, clear=args.clear)
    except FileNotFoundError as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    _cli()
